from __future__ import annotations

import argparse
from datetime import datetime, timedelta
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
from PIL import Image

from project_config import CONFIG


PATHS_CONFIG = CONFIG["paths"]
IMAGES_CONFIG = CONFIG["images"]
REPORT_CONFIG = CONFIG["report"]
SUPPORTED_EXTENSIONS = {
    extension.lower() for extension in IMAGES_CONFIG["supported_extensions"]
}
MAX_IMAGE_WIDTH_CM = REPORT_CONFIG["max_image_width_cm"]
MAX_IMAGE_HEIGHT_CM = REPORT_CONFIG["max_image_height_cm"]
PAGE_WIDTH_CM = REPORT_CONFIG["page_width_cm"]
PAGE_HEIGHT_CM = REPORT_CONFIG["page_height_cm"]
DEFAULT_TEMPLATE = PATHS_CONFIG["template_file"]
DEFAULT_TITLE = REPORT_CONFIG["title"]
PICS_DIR = PATHS_CONFIG["pics_dir"]
COLLAGES_DIR = PATHS_CONFIG["collages_dir"]
FILENAME_PREFIX = REPORT_CONFIG["filename_prefix"]
FOLDER_DATE_FORMAT = REPORT_CONFIG["folder_date_format"]
FILENAME_DATE_FORMAT = REPORT_CONFIG["filename_date_format"]
HEADING_MONTH_YEAR_FORMAT = REPORT_CONFIG["heading_month_year_format"]


def format_folder_heading(name: str) -> str:
    try:
        date_value = datetime.strptime(name, FOLDER_DATE_FORMAT)
    except ValueError:
        return name
    return f"{date_value.day} {date_value.strftime(HEADING_MONTH_YEAR_FORMAT)}"


def report_directories(root: Path) -> list[Path]:
    return sorted(
        [path for path in root.iterdir() if path.is_dir()],
        key=lambda path: path.name,
    )


def image_files(root: Path) -> list[Path]:
    return sorted(
        [
            path
            for path in root.iterdir()
            if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS
        ],
        key=lambda path: path.name.lower(),
    )


def add_page_break(document: Document) -> None:
    document.add_page_break()


def resolve_case_insensitive(path_str: str) -> Path:
    path = Path(path_str).expanduser()
    if not path.is_absolute():
        path = Path.cwd() / path

    path = path.resolve(strict=False)
    if path.exists():
        return path

    anchor = Path(path.anchor) if path.anchor else Path.cwd().anchor
    current = Path(anchor) if isinstance(anchor, str) else anchor

    for part in path.parts[len(current.parts) :]:
        if not current.exists() or not current.is_dir():
            return path

        match = next(
            (child for child in current.iterdir() if child.name.lower() == part.lower()),
            None,
        )
        if match is None:
            return path
        current = match

    return current


def report_filename_suffix_from_pics() -> str:
    pics_root = resolve_case_insensitive(PICS_DIR)
    if not pics_root.exists():
        raise SystemExit(f"Folder not found: {pics_root}")

    folder_names = sorted(
        [path.name for path in pics_root.iterdir() if path.is_dir()],
        key=str.lower,
    )
    if not folder_names:
        raise SystemExit(f"No folders found in {pics_root}")

    date_values = []
    for name in folder_names:
        try:
            date_values.append(datetime.strptime(name, FOLDER_DATE_FORMAT))
        except ValueError:
            continue

    if date_values:
        earliest = min(date_values)
        monday = earliest - timedelta(days=earliest.weekday())
        return monday.strftime(FILENAME_DATE_FORMAT)

    return datetime.now().strftime(FILENAME_DATE_FORMAT)


def default_output_filename() -> str:
    suffix = report_filename_suffix_from_pics()
    return f"{FILENAME_PREFIX}{suffix}.docx"


def configure_page_size(document: Document) -> None:
    for section in document.sections:
        section.page_width = Cm(PAGE_WIDTH_CM)
        section.page_height = Cm(PAGE_HEIGHT_CM)


def add_page_numbers(document: Document) -> None:
    for section in document.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        run = paragraph.add_run()
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")

        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = "PAGE"

        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")

        run._r.append(fld_char_begin)
        run._r.append(instr_text)
        run._r.append(fld_char_end)


def picture_size(picture: Path) -> tuple[Cm, Cm]:
    with Image.open(picture) as image:
        width_px, height_px = image.size

    scale = min(MAX_IMAGE_WIDTH_CM / width_px, MAX_IMAGE_HEIGHT_CM / height_px)
    width_cm = width_px * scale
    height_cm = height_px * scale
    return Cm(width_cm), Cm(height_cm)


def build_report(
    source_root: Path,
    output_path: Path,
    template_path: Path | None,
    title: str,
) -> None:
    document = Document(str(template_path)) if template_path else Document()
    configure_page_size(document)
    add_page_numbers(document)
    document.add_heading(title, level=0)
    first_section = True
    first_picture_in_section = True

    for report_dir in report_directories(source_root):
        pictures = image_files(report_dir)
        if not pictures:
            continue

        if not first_section:
            add_page_break(document)
        first_section = False
        first_picture_in_section = True

        document.add_heading(format_folder_heading(report_dir.name), level=1)

        for picture in pictures:
            if not first_picture_in_section:
                add_page_break(document)
            first_picture_in_section = False

            document.add_heading(picture.stem.upper(), level=2)

            width, height = picture_size(picture)
            document.add_picture(str(picture), width=width, height=height)

    document.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(
        description=f"Generate a Word report from pictures in {COLLAGES_DIR}."
    )
    parser.add_argument(
        "source",
        nargs="?",
        default=COLLAGES_DIR,
        help=f"Root folder containing report folders. Defaults to ./{COLLAGES_DIR}",
    )
    parser.add_argument(
        "--output",
        help="Optional output .docx filename. Defaults to a name derived from the title.",
    )
    parser.add_argument(
        "--title",
        default=DEFAULT_TITLE,
        help=f"Report title shown at the start of the document. Defaults to {DEFAULT_TITLE}.",
    )
    parser.add_argument(
        "--template",
        default=DEFAULT_TEMPLATE,
        help=(
            "Optional .docx template to start from. "
            "Use a template with the Gallery theme already applied."
        ),
    )
    args = parser.parse_args()

    if MAX_IMAGE_WIDTH_CM <= 0 or MAX_IMAGE_HEIGHT_CM <= 0:
        raise SystemExit("report maximum image dimensions in config.toml must be positive")
    if PAGE_WIDTH_CM <= 0 or PAGE_HEIGHT_CM <= 0:
        raise SystemExit("report page dimensions in config.toml must be positive")

    source_root = Path(args.source).resolve()
    if not source_root.exists():
        raise SystemExit(f"Folder not found: {source_root}")

    output_name = args.output or default_output_filename()
    output_path = Path(output_name).resolve()
    template_path = Path(args.template).resolve()
    if not template_path.exists():
        template_path = None

    build_report(source_root, output_path, template_path, args.title)
    print(f"Created report: {output_path}")


if __name__ == "__main__":
    main()
